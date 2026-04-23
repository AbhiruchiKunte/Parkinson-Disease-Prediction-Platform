
try:
    from docx import Document
except ImportError:
    Document = None

try:
    from pypdf import PdfReader
except ImportError:
    PdfReader = None

import pandas as pd
import io
import re


def _normalize_column_name(name):
    raw = str(name).strip().lower()
    raw = re.sub(r'[^a-z0-9]+', '_', raw).strip('_')

    alias_map = {
        'patientage': 'age',
        'subject_age': 'age',
        'years_old': 'age',
        'tremor': 'tremor_score',
        'rest_tremor': 'tremor_score',
        'handwriting': 'handwriting_score',
        'micrographia': 'handwriting_score',
        'jitter': 'jitter_local',
        'jitter_percent': 'jitter_local',
        'jitter_perc': 'jitter_local',
        'mdvp_jitter': 'jitter_local',
        'shimmer': 'shimmer_local',
        'shimmer_db': 'shimmer_local',
        'mdvp_shimmer': 'shimmer_local',
    }
    return alias_map.get(raw, raw)


def _normalize_dataframe(df):
    if df is None or df.empty:
        return df
    df.columns = [_normalize_column_name(c) for c in df.columns]
    return df


def _extract_key_value_records(text):
    """
    Extract multiple patient-like records from free text blocks.
    """
    patterns = {
        'age': [r'age\s*[:=-]?\s*(\d{1,3})', r'(\d{1,3})\s*years?\s*old'],
        'tremor_score': [r'tremor(?:\s*score)?\s*[:=-]?\s*(\d+(?:\.\d+)?)'],
        'handwriting_score': [r'handwriting(?:\s*score)?\s*[:=-]?\s*(\d+(?:\.\d+)?)', r'micrographia\s*[:=-]?\s*(\d+(?:\.\d+)?)'],
        'jitter_local': [r'jitter(?:\s*local)?(?:\s*\(abs\)|\s*%)?\s*[:=-]?\s*(\d+(?:\.\d+)?)'],
        'shimmer_local': [r'shimmer(?:\s*local)?(?:\s*\(db\)|\s*%)?\s*[:=-]?\s*(\d+(?:\.\d+)?)'],
        'bradykinesia': [r'bradykinesia\s*[:=-]?\s*(\d+(?:\.\d+)?)'],
        'rigidity': [r'rigidity\s*[:=-]?\s*(\d+(?:\.\d+)?)'],
    }

    blocks = re.split(
        r'\n\s*\n+|(?=patient\s*(?:id|no|number)?\s*[:#-]?\s*[a-z0-9_-]+)',
        text,
        flags=re.IGNORECASE
    )

    records = []
    for block in blocks:
        block = block.strip()
        if not block:
            continue
        row = {}
        for field, regex_list in patterns.items():
            for regex in regex_list:
                match = re.search(regex, block, re.IGNORECASE)
                if match:
                    row[field] = match.group(1)
                    break
        # Keep only meaningful rows
        if len(row) >= 3:
            records.append(row)

    if not records:
        # Last chance: parse entire text as one record if it contains enough fields
        row = {}
        for field, regex_list in patterns.items():
            for regex in regex_list:
                match = re.search(regex, text, re.IGNORECASE)
                if match:
                    row[field] = match.group(1)
                    break
        if len(row) >= 3:
            records.append(row)

    return pd.DataFrame(records) if records else pd.DataFrame()


def _extract_pd_table_rows(text):
    """
    Fallback extractor for PDF/DOC text where table headers are broken but
    data rows are still readable, e.g.:
    PD_001 0.42 0.71 0.56 4.8 0.59 18 Moderate
    """
    lines = [ln.strip() for ln in text.replace('\r', '\n').split('\n') if ln.strip()]
    row_pattern = re.compile(
        r'^(?P<patient>[a-z]{1,4}[_\-]?\d{2,6})\s+'
        r'(?P<spiral_error>-?\d+(?:\.\d+)?)\s+'
        r'(?P<drawing_pressure>-?\d+(?:\.\d+)?)\s+'
        r'(?P<stroke_velocity>-?\d+(?:\.\d+)?)\s+'
        r'(?P<tremor_frequency>-?\d+(?:\.\d+)?)\s+'
        r'(?P<writing_speed>-?\d+(?:\.\d+)?)\s+'
        r'(?P<pen_lift_count>-?\d+(?:\.\d+)?)\s+'
        r'(?P<handwriting_score>[a-zA-Z]+)$',
        re.IGNORECASE,
    )

    records = []
    for line in lines:
        match = row_pattern.match(line)
        if not match:
            continue
        row = match.groupdict()
        row["patient_id"] = row.pop("patient")
        records.append(row)

    if records:
        return _normalize_dataframe(pd.DataFrame(records))
    return pd.DataFrame()


def parse_text_table(text):
    """
    Attempts to parse text containing a table (whitespace/tab separated).
    Returns DataFrame or None if no table structure found.
    Handles pandas read_csv failures by falling back to manual line parsing.
    """
    # Pre-process text to remove common garbage or empty lines
    lines = [line.strip() for line in text.replace('\r', '\n').split('\n') if line.strip()]
    if len(lines) < 2:
        return None

    clean_text = '\n'.join(lines)

    # Strategy 1: delimiter inference across common table separators
    separators = [',', ';', '\t', '|', r'\s{2,}']
    for sep in separators:
        try:
            df = pd.read_csv(io.StringIO(clean_text), sep=sep, engine='python')
            if len(df) > 0 and len(df.columns) > 1:
                return _normalize_dataframe(df)
        except Exception:
            continue

    # Strategy 2: Manual Line Parsing (Robust Fallback)
    try:
        data = []
        headers = []
        header_idx = -1

        # Keywords to identify likely header row
        keywords = ['patient', 'id', 'age', 'tremor', 'handwriting', 'spiral', 'drawing', 'pressure', 'velocity', 'speed', 'jitter', 'shimmer', 'score', 'count']

        # Find header line
        for i, line in enumerate(lines):
            lower_line = line.lower()
            matches = sum(1 for k in keywords if k in lower_line)
            if matches >= 2 and len(re.split(r'[,\t;|]|\s{2,}', line)) >= 2:
                headers = [h.strip() for h in re.split(r'[,\t;|]|\s{2,}', line) if h.strip()]
                header_idx = i
                break

        if header_idx != -1 and headers:
            for line in lines[header_idx+1:]:
                parts = [p.strip() for p in re.split(r'[,\t;|]|\s{2,}', line) if p.strip()]
                if len(parts) >= len(headers):
                    data.append(dict(zip(headers, parts[:len(headers)])))

            if data:
                return _normalize_dataframe(pd.DataFrame(data))

    except Exception as e:
        print(f"Manual parsing failed: {e}")
        pass

    return None

def parse_docx_to_df(file_stream):
    """Parse a .docx file and return a DataFrame."""
    if not Document:
        raise ImportError("python-docx is not installed.")
    
    file_stream.seek(0)
    document = Document(file_stream)
    data = []
    
    # Strategy 1: Look for tables (most reliable for structured data)
    for table in document.tables:
        headers = []
        for i, row in enumerate(table.rows):
            row_data = [cell.text.strip() for cell in row.cells]
            if i == 0:
                headers = [_normalize_column_name(h) for h in row_data]
            else:
                # Create dict based on headers
                if len(row_data) == len(headers):
                    data.append(dict(zip(headers, row_data)))
    
    # If tables found, return DF
    if data:
        return _normalize_dataframe(pd.DataFrame(data))
        
    # Strategy 2: Look for tabular text in paragraphs (fallback for "Text Tables")
    full_text = '\n'.join([p.text for p in document.paragraphs])
    df_table = parse_text_table(full_text)
    if df_table is not None:
        return _normalize_dataframe(df_table)

    # Strategy 3: Multiple key/value records
    kv_df = _extract_key_value_records(full_text)
    if not kv_df.empty:
        return _normalize_dataframe(kv_df)

    # Strategy 4: Single Key-Value pairs (fallback)
    row_dict = {}
    for para in document.paragraphs:
        text = para.text.strip()
        if ':' in text:
            parts = text.split(':', 1)
            key = _normalize_column_name(parts[0].strip())
            val = parts[1].strip()
            row_dict[key] = val
    
    if row_dict:
        return _normalize_dataframe(pd.DataFrame([row_dict]))
        
    return pd.DataFrame() # Empty if nothing found


def parse_doc_to_df(file_stream):
    """
    Attempt to parse legacy .doc files.
    """
    # 1. Try as docx first
    try:
        if Document:
            file_stream.seek(0)
            return parse_docx_to_df(file_stream)
    except Exception:
        pass # Not a valid docx
        
    # 2. Binary string extraction (Best effort for .doc)
    try:
        file_stream.seek(0)
        content = file_stream.read()
        # Decode with ignore to get printable strings
        text = content.decode('latin-1', errors='ignore')
        
        # Strategy A: Text Table
        df_table = parse_text_table(text)
        if df_table is not None:
            return _normalize_dataframe(df_table)

        # Strategy B: Row-pattern extraction from flattened table text
        row_df = _extract_pd_table_rows(text)
        if not row_df.empty:
            return _normalize_dataframe(row_df)

        # Strategy C: Multi-record key/value extraction
        kv_df = _extract_key_value_records(text)
        if not kv_df.empty:
            return _normalize_dataframe(kv_df)
            
        # Strategy D: Key-Value Regex (Legacy Fallback)
        # Look for keywords in the raw text
        row_dict = {}
        
        # Define patterns for our fields
        patterns = {
            'age': [r'age.*?(\d+)', r'years.*?old.*?(\d+)'],
            'tremor_score': [r'tremor.*?(\d+(\.\d+)?)'],
            'handwriting_score': [r'handwriting.*?(\d+(\.\d+)?)', r'micrographia.*?(\d+(\.\d+)?)'],
            'jitter_local': [r'jitter.*?(\d+(\.\d+)?)'],
            'shimmer_local': [r'shimmer.*?(\d+(\.\d+)?)'],
            'bradykinesia': [r'bradykinesia.*?(\d+(\.\d+)?)'],
            'rigidity': [r'rigidity.*?(\d+(\.\d+)?)']
        }
        
        for field, regex_list in patterns.items():
            for regex in regex_list:
                match = re.search(regex, text, re.IGNORECASE)
                if match:
                    row_dict[field] = match.group(1)
                    break
        
        if row_dict:
            return _normalize_dataframe(pd.DataFrame([row_dict]))
            
    except Exception as e:
        print(f"Doc parsing error: {e}")
        pass
        
    return pd.DataFrame()

def parse_pdf_to_df(file_stream):
    """Parse a .pdf file and return a DataFrame."""
    if not PdfReader:
        raise ImportError("pypdf is not installed.")
        
    reader = PdfReader(file_stream)
    text = ""
    for page in reader.pages:
        page_text = page.extract_text() or ""
        text += page_text + "\n"
    
    # Strategy 1: Text Table (Generic whitespace separated)
    df_table = parse_text_table(text)
    if df_table is not None:
        return _normalize_dataframe(df_table)

    # Strategy 2: Row-pattern extraction from flattened table text
    row_df = _extract_pd_table_rows(text)
    if not row_df.empty:
        return _normalize_dataframe(row_df)

    # Strategy 3: Multi-record key/value extraction
    kv_df = _extract_key_value_records(text)
    if not kv_df.empty:
        return _normalize_dataframe(kv_df)

    # Strategy 4: Robust Key-Value Parsing (Single-record fallback)
    row_dict = {}
    
    # Common patterns for medical reports
    # Age: 60, Age 60, 60 years old
    # Tremor: 1, Tremor Score: 1
    
    patterns = {
        'age': [r'age\s*[:=-]?\s*(\d+)', r'(\d+)\s*years\s*old'],
        'tremor_score': [r'tremor\s*(?:score)?\s*[:=-]?\s*(\d+(?:\.\d+)?)'],
        'handwriting_score': [r'handwriting\s*(?:score)?\s*[:=-]?\s*(\d+(?:\.\d+)?)', r'micrographia\s*[:=-]?\s*(\d+(?:\.\d+)?)'],
        'jitter_local': [r'jitter\s*(?:local)?\s*(?:\(abs\))?\s*[:=-]?\s*(\d+(?:\.\d+)?)'],
        'shimmer_local': [r'shimmer\s*(?:local)?\s*[:=-]?\s*(\d+(?:\.\d+)?)'],
        'bradykinesia': [r'bradykinesia\s*[:=-]?\s*(\d+(?:\.\d+)?)'],
        'rigidity': [r'rigidity\s*[:=-]?\s*(\d+(?:\.\d+)?)']
    }
    
    for field, regex_list in patterns.items():
        found = False
        for regex in regex_list:
            match = re.search(regex, text, re.IGNORECASE)
            if match:
                row_dict[field] = match.group(1)
                found = True
                break
                
    if row_dict:
        return _normalize_dataframe(pd.DataFrame([row_dict]))
        
    return pd.DataFrame()

# Helper function for batch processing
def process_csv_batch(df, model_handler):
    results = []
    failed = 0
    successful = 0
    
    # Normalize column names for easier matching
    df.columns = [str(c).lower().strip() for c in df.columns]
    
    def find_column(keywords, row):
        """Find value for a feature by searching for keywords in column names"""
        # 1. Try exact matches first
        for key in keywords:
            if key in row:
                val = row[key]
                # Clean value if it's a string
                if isinstance(val, str):
                   import re
                   val_clean = re.sub(r'[^\d.]', '', val)
                   try:
                       return float(val_clean)
                   except:
                       pass 
                return val
        
        # 2. Try partial matches (fuzzy)
        for col in row.index:
            for key in keywords:
                if key in col:
                     val = row[col]
                     if isinstance(val, str):
                        import re
                        val_clean = re.sub(r'[^\d.]', '', val)
                        try:
                            return float(val_clean)
                        except:
                            pass
                     return val
        return 0 # Default if not found

    for index, row in df.iterrows():
        try:
            # INTELLIGENT MAPPING: Expanded keywords based on user dataset
            jitter = find_column(['jitter_local', 'jitter_perc', 'jitter_abs', 'jitter', 'jit'], row)
            shimmer = find_column(['shimmer_local', 'shimmer_apq3', 'shimmer_apq5', 'shimmer', 'shim'], row)
            tremor = find_column(['tremor_score', 'tremor', 'tremor_frequency', 'tremor_f'], row)
            age = find_column(['age', 'patient_age', 'years'], row) or 60 
            
            # Map other user columns
            handwriting = find_column(['handwriting_score', 'handwriting', 'micrographia', 'handwriting_s'], row)
            bradykinesia = find_column(['bradykinesia', 'slowness', 'stroke_velocity', 'writing_speed', 'stroke_v', 'writing_s'], row)
            rigidity = find_column(['rigidity', 'stiffness', 'drawing_pressure', 'drawing_p', 'pen_lift'], row)

            # Ensure safe conversion to float
            def safe_float(v, default=0.0):
                try:
                    return float(v)
                except:
                    return default

            features = {
                'age': safe_float(age, 60),
                'tremor_score': safe_float(tremor),
                'handwriting_score': safe_float(handwriting),
                'jitter_local': safe_float(jitter),
                'shimmer_local': safe_float(shimmer),
                'bradykinesia': safe_float(bradykinesia),
                'rigidity': safe_float(rigidity)
            }
            
            # Run prediction
            pred = model_handler.predict_single(features)
            
            results.append({
                "row": index + 1,
                "input": features,
                "prediction": pred,
                "status": "success"
            })
            successful += 1
            
        except Exception as e:
            results.append({
                "row": index + 1,
                "error": str(e),
                "status": "failed"
            })
            failed += 1
            
    return {
        "predictions": results,
        "total_records": len(df),
        "successful_predictions": successful,
        "failed_predictions": failed
    }
